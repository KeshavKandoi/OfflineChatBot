from uuid import uuid4
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.document_loaders import PyPDFLoader
from langchain_ollama import ChatOllama
from langchain_core.messages import HumanMessage
import docx
import base64


# Embedding model
embedding_model = HuggingFaceEmbeddings(
    model_name="sentence-transformers/all-MiniLM-L6-v2"
)


# Separate ChromaDB collection for documents
rag_store = Chroma(
    collection_name="rag_documents",
    embedding_function=embedding_model,
    persist_directory="./chroma_db",
)


# Vision model for images
vision_model = ChatOllama(model="minicpm-v")


# Text splitter
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=500,
    chunk_overlap=50,
)


# ── PDF ──
def process_pdf(file_path: str, filename: str):
    loader = PyPDFLoader(file_path)
    docs = loader.load()
    chunks = text_splitter.split_documents(docs)

    final_docs = []

    for i, chunk in enumerate(chunks):
        doc = Document(
            page_content=chunk.page_content,
            metadata={
                "filename": filename,
                "type": "pdf",
                "chunk_index": i
            }
        )
        final_docs.append(doc)

    uuids = [str(uuid4()) for _ in final_docs]
    rag_store.add_documents(documents=final_docs, ids=uuids)

    return len(final_docs)


# ── TXT ──
def process_txt(file_path: str, filename: str):
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()

    chunks = text_splitter.split_text(text)

    final_docs = []

    for i, chunk in enumerate(chunks):
        doc = Document(
            page_content=chunk,
            metadata={
                "filename": filename,
                "type": "txt",
                "chunk_index": i
            }
        )
        final_docs.append(doc)

    uuids = [str(uuid4()) for _ in final_docs]
    rag_store.add_documents(documents=final_docs, ids=uuids)

    return len(final_docs)


# ── DOCX ──
def process_docx(file_path: str, filename: str):
    doc = docx.Document(file_path)

    texts = []

    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)

    text = "\n".join(texts)

    chunks = text_splitter.split_text(text)

    final_docs = []

    for i, chunk in enumerate(chunks):
        doc = Document(
            page_content=chunk,
            metadata={
                "filename": filename,
                "type": "docx",
                "chunk_index": i
            }
        )
        final_docs.append(doc)

    uuids = [str(uuid4()) for _ in final_docs]
    rag_store.add_documents(documents=final_docs, ids=uuids)

    return len(final_docs)


# ── IMAGE ──
def process_image(file_path: str, filename: str):
    ext = file_path.lower().split(".")[-1]
    mime = "image/png" if ext == "png" else "image/jpeg" if ext in ["jpg", "jpeg"] else "image/png"
    with open(file_path, "rb") as f:
        image_data = base64.b64encode(f.read()).decode("utf-8")

    message = HumanMessage(content=[
       {"type": "text", "text": (
            "Analyze the actual content of this image in detail. "
            "If it contains a diagram, chart, architecture, code, or technical content, "
            "explain what it shows — components, labels, connections, flow, and purpose. "
            "If it contains text, extract and explain it. "
            "Do NOT describe file metadata, screenshot borders, desktop, or timestamps. "
            "Focus entirely on the subject matter inside the image."
        )},
        {"type": "image_url", "image_url": f"data:{mime};base64,{image_data}"}
    ])

    response = vision_model.invoke([message])
    description = response.content

    doc = Document(
        page_content=description,
        metadata={
            "filename": filename,
            "type": "image"
        }
    )

    rag_store.add_documents(documents=[doc], ids=[str(uuid4())])

    return description


# ── SEARCH ──
def search_rag(query: str, n_results: int = 6):
    try:
        results = rag_store.similarity_search(query, k=n_results)

        contents = []
        for doc in results:
            contents.append(doc.page_content)

        return contents

    except Exception:
        return []